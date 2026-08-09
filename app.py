#!/usr/bin/env python3
"""
app.py - Branch Activation Email Reviewer

Reads AOsheet.xlsx (account activation log) and EmailIDs.xlsx (branch -> email
lookup), groups AOsheet rows by Branch Code (Column H), and opens a colorful
review GUI that shows one branch's email at a time -- editable To, Cc, and
Subject fields, plus a read-only preview of the email body (the account
table + fixed wording; the body itself cannot be edited) -- with Send /
Cancel Sending buttons at the bottom, and an End button (side) to stop the
whole run early at any point.

  - Send  -> sends that email now (via the SMIFS SMTP relay) exactly as shown
             on screen, then moves to the next one.
  - Cancel Sending -> skips that email (nothing sent for it), moves to the next.
  - End   -> stops the review immediately; any branches not yet reached are
             recorded as "not sent".

Once every branch has been handled (or End was pressed), the GUI shows a
"Process Completed" summary (sent count / not-sent count), and after you
close that, a Word document is generated automatically (no save dialog) --
saved in the same folder as AOsheet.xlsx, filename suffixed with today's
date -- listing every email that was sent (with its To/Cc/Subject and the
actual email text/table as sent) and, separately, every branch that was
not sent, in sequence.

The email wording/format (greeting, N.B. note, helpdesk line, signature) is
fixed in this script (see EMAIL_* constants below) and is NOT editable in the
GUI -- only shown as a preview. Emails are sent as HTML (with a plain-text
fallback auto-generated from the same source) so the account table renders
properly in the recipient's inbox.
The SMTP password is hardcoded below (SMTP_PASSWORD) at the user's request --
see the WARNING comment above that constant regarding the security tradeoff.

Usage:
    python app.py

Requires: openpyxl, python-docx   (pip install openpyxl python-docx)
Optional: tkinterweb -- if installed, a live rendered preview pane appears
next to the HTML editor (pip install tkinterweb). Without it, the app still
works fully (edit + send), just without the visual preview pane.
Tkinter ships with standard Python on Windows/Mac; on Linux you may need
to install it separately (e.g. `sudo apt install python3-tk`).
"""

import sys
import os
import socket
import smtplib
from email.message import EmailMessage
from email.utils import formatdate, make_msgid
from collections import OrderedDict
from datetime import datetime

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

# --------------------------------------------------------------------------
# FIXED EMAIL FORMAT -- edit here if the wording ever needs to change.
# This is intentionally hardcoded so the script never has to ask for it.
# --------------------------------------------------------------------------

EMAIL_GREETING = "Dear Team,"
EMAIL_INTRO_1 = "Greetings from SMIFS Limited!"
EMAIL_INTRO_2 = (
    "We are glad to inform you that the below-mentioned Accounts logged "
    "from your Branch have been activated today:"
)
EMAIL_NB = (
    "N.B. - Please note that the Clients can transact in the Securities "
    "Market subject to receiving Approval to Trade from respective "
    "Exchanges and completion of UCC Mapping by the Depositories."
)
EMAIL_HELPDESK = (
    "In case you have any query, you may reach our Helpdesk Team on "
    "+91 33 4057 2625 / 35 / 45 or write to them at helpdesk@smifs.com."
)
EMAIL_SIGNATURE_LINES = [
    "Regards,",
    "AMIT KUMAR SHAW",
    "Chief Manager Operations",
    "Mackertich ONE",
    "SMIFS Limited",
    "5F Vaibhav, 4 Lee Road, Kolkata-700020",
    "Phone: 9903300266 / 9903238970",
    "www.smifs.com",
]
SUBJECT_TEMPLATE = "Account Activation Confirmation — {branch_name} ({branch_code})"
DEFAULT_CC = "account.opening@smifs.com"   # pre-filled in Cc for every email; editable per email

# --------------------------------------------------------------------------
# SMTP CONFIG -- emails are sent (not drafted) via this relay.
# WARNING: the password below is hardcoded at the user's explicit request.
# This file is now sensitive -- do not share, email, or commit it anywhere
# it could be seen by others. Rotate the password if that ever happens.
# --------------------------------------------------------------------------

SENDER_EMAIL = "kallol.saha@smifs.com"   # used as the From: address

SMTP_SERVER = "180.179.151.1"
SMTP_PORT = 587
SMTP_USERNAME = "hrm-smifspr@m3c.in"     # relay account used to authenticate
SMTP_PASSWORD = "Sm7eUkY6W9IiT"          # hardcoded per user request -- see WARNING above

# --------------------------------------------------------------------------
# COLUMN MAP -- AOsheet.xlsx column letter -> output heading (12 columns)
# Row 1/2 of AOsheet.xlsx are header rows; data starts at row 3.
# --------------------------------------------------------------------------

COLUMN_MAP = [
    ("B", "Active Date"),
    ("C", "UCC"),
    ("D", "Client Name"),
    ("E", "Client Category"),
    ("H", "Branch Code"),
    ("I", "Branch Name"),
    ("J", "AOF Type"),
    ("K", "DP ID"),
    ("L", "Client ID"),
    ("O", "RM Name"),
    ("P", "DDPI"),
    ("Q", "Remarks"),
]

BRANCH_CODE_COL = "H"          # branch code column used for grouping
BRANCH_NAME_COL = "I"          # branch name column, used for subject/index

# EmailIDs.xlsx layout: Column A = branch code, Column C = email address
EMAILIDS_BRCODE_COL = "A"
EMAILIDS_EMAIL_COL = "C"

DATA_START_ROW = 3   # AOsheet.xlsx: row 1 = grouped header, row 2 = sub-header
EMAILIDS_START_ROW = 2  # EmailIDs.xlsx: row 1 = header


def col_letter_to_index(letter):
    """Convert an Excel column letter (e.g. 'H') to a 1-based column index."""
    idx = 0
    for ch in letter:
        idx = idx * 26 + (ord(ch.upper()) - ord('A') + 1)
    return idx


def fmt_value(value):
    """Format a cell value for display in the email table."""
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d-%m-%Y")
    return str(value)


def read_ao_sheet(path):
    """
    Read AOsheet.xlsx and group rows by Branch Code (column H).
    Returns an OrderedDict: { branch_code: {"branch_name": str, "rows": [dict, ...]} }
    preserving the order branch codes first appear in the sheet.
    """
    wb = load_workbook(path, data_only=True)
    ws = wb.active

    col_indices = {letter: col_letter_to_index(letter) for letter, _ in COLUMN_MAP}
    branch_code_idx = col_letter_to_index(BRANCH_CODE_COL)
    branch_name_idx = col_letter_to_index(BRANCH_NAME_COL)

    groups = OrderedDict()
    for row in ws.iter_rows(min_row=DATA_START_ROW):
        branch_code = row[branch_code_idx - 1].value
        if branch_code is None or str(branch_code).strip() == "":
            continue
        branch_code = str(branch_code).strip()
        branch_name = row[branch_name_idx - 1].value or ""

        record = {}
        for letter, heading in COLUMN_MAP:
            cell = row[col_indices[letter] - 1]
            record[heading] = fmt_value(cell.value)

        if branch_code not in groups:
            groups[branch_code] = {"branch_name": branch_name, "rows": []}
        groups[branch_code]["rows"].append(record)

    return groups


def read_email_map(path):
    """
    Read EmailIDs.xlsx and build { branch_code: email_address } using
    Column A (branch code) and Column C (email id).
    """
    wb = load_workbook(path, data_only=True)
    ws = wb.active

    brcode_idx = col_letter_to_index(EMAILIDS_BRCODE_COL)
    email_idx = col_letter_to_index(EMAILIDS_EMAIL_COL)

    email_map = {}
    for row in ws.iter_rows(min_row=EMAILIDS_START_ROW):
        code_cell = row[brcode_idx - 1].value
        email_cell = row[email_idx - 1].value
        if code_cell is None:
            continue
        code = str(code_cell).strip()
        if code and email_cell:
            email_map[code] = str(email_cell).strip()

    return email_map


def build_html_table(rows):
    """Render the account rows as an HTML <table> (used in the editable
    email body, and rendered live in the preview pane)."""
    headings = [h for _, h in COLUMN_MAP]
    thead = "".join(
        f'<th style="border:1px solid #999;padding:4px 6px;background:#d9d9d9;'
        f'font-family:Calibri,Arial,sans-serif;font-size:12px;text-align:left;">{h}</th>'
        for h in headings
    )
    body_rows = ""
    for record in rows:
        cells = "".join(
            f'<td style="border:1px solid #999;padding:4px 6px;'
            f'font-family:Calibri,Arial,sans-serif;font-size:12px;">{record.get(h, "")}</td>'
            for h in headings
        )
        body_rows += f"<tr>{cells}</tr>"
    return (
        f'<table style="border-collapse:collapse;border:1px solid #999;">'
        f"<tr>{thead}</tr>{body_rows}</table>"
    )


def build_default_body(rows):
    """Default (pre-fill) HTML body for one branch's email."""
    table_html = build_html_table(rows)
    signature_html = "<br>".join(EMAIL_SIGNATURE_LINES)
    html = f"""<html><body style="font-family:Calibri,Arial,sans-serif;font-size:14px;">
<p>{EMAIL_GREETING}</p>
<p>{EMAIL_INTRO_1}</p>
<p>{EMAIL_INTRO_2}</p>
{table_html}
<p><b>{EMAIL_NB}</b></p>
<p>{EMAIL_HELPDESK}</p>
<p>{signature_html}</p>
</body></html>"""
    return html


def html_to_plain_fallback(html):
    """Very simple HTML -> plain-text fallback for the multipart/alternative
    plain part (most mail clients will show the HTML part instead)."""
    import re
    text = re.sub(r"(?i)<br\s*/?>", "\n", html)
    text = re.sub(r"(?i)</(p|tr|div|h[1-6])>", "\n", text)
    text = re.sub(r"(?i)<td[^>]*>", " | ", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()


def connect_smtp():
    """Open and authenticate an SMTP connection. Raises on failure."""
    socket.setdefaulttimeout(20)
    conn = smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=20)
    conn.ehlo()
    if conn.has_extn("STARTTLS"):
        conn.starttls()
        conn.ehlo()
    conn.login(SMTP_USERNAME, SMTP_PASSWORD)
    return conn


def parse_address_list(raw):
    """Split a To/Cc field on comma or semicolon into a clean list."""
    raw = raw.replace(";", ",")
    return [a.strip() for a in raw.split(",") if a.strip()]


# ==========================================================================
# GUI -- one email at a time: editable To / Cc / Subject (body is fixed,
# read-only), with Send / Cancel Sending / End controls.
# ==========================================================================

# Color palette used throughout the GUI.
COLOR_BG = "#eef3f8"
COLOR_HEADER_BG = "#2c3e50"
COLOR_HEADER_FG = "#ffffff"
COLOR_ACCENT = "#2980b9"
COLOR_CARD_BG = "#ffffff"
COLOR_LABEL_FG = "#2c3e50"
COLOR_SEND_BG = "#27ae60"
COLOR_CANCEL_BG = "#e67e22"
COLOR_END_BG = "#c0392b"
COLOR_PROGRESS_BG = "#3498db"


class EmailReviewApp:
    def __init__(self, root, branches, email_map, sender_email):
        self.root = root
        self.branches = branches            # ordered list of dicts: branch_code, branch_name, rows
        self.email_map = email_map
        self.sender_email = sender_email
        self.index = 0
        self.current_body = ""              # fixed (non-editable) HTML body for the current branch
        self.results = []                   # filled in as user sends/cancels each email
        self.sent_count = 0
        self.not_sent_count = 0
        self.smtp_conn = None

        self.root.title("Branch Activation Email Review")
        self.root.geometry("1150x850")
        self.root.configure(bg=COLOR_BG)
        self._build_widgets()
        self._load_current()

    # -- layout -----------------------------------------------------------
    def _build_widgets(self):
        import tkinter as tk
        try:
            from tkinterweb import HtmlFrame
            self.preview_available = True
        except ImportError:
            self.preview_available = False

        # --- Header banner ---
        header = tk.Frame(self.root, bg=COLOR_HEADER_BG, height=56)
        header.pack(fill="x", side="top")
        tk.Label(
            header, text="\U0001F4E7  Branch Activation Email Review",
            bg=COLOR_HEADER_BG, fg=COLOR_HEADER_FG,
            font=("Segoe UI", 16, "bold"), pady=12,
        ).pack(side="left", padx=16)

        outer = tk.Frame(self.root, bg=COLOR_BG)
        outer.pack(fill="both", expand=True, padx=14, pady=14)

        main = tk.Frame(outer, bg=COLOR_BG)
        main.pack(side="left", fill="both", expand=True)

        side = tk.Frame(outer, width=100, bg=COLOR_BG)
        side.pack(side="right", fill="y", padx=(14, 0))

        # --- Progress badge ---
        self.progress_label = tk.Label(
            main, text="", font=("Segoe UI", 11, "bold"),
            bg=COLOR_PROGRESS_BG, fg="white", anchor="w", padx=12, pady=6,
        )
        self.progress_label.pack(fill="x", pady=(0, 10))

        # --- Form card (To / Cc / Subject) ---
        form_card = tk.Frame(main, bg=COLOR_CARD_BG, highlightbackground=COLOR_ACCENT,
                              highlightthickness=1, padx=14, pady=12)
        form_card.pack(fill="x", pady=(0, 10))

        entry_style = dict(font=("Segoe UI", 10), relief="solid", borderwidth=1,
                            highlightthickness=0)
        label_style = dict(bg=COLOR_CARD_BG, fg=COLOR_LABEL_FG, font=("Segoe UI", 10, "bold"))

        tk.Label(form_card, text="To:", width=8, anchor="w", **label_style).grid(
            row=0, column=0, sticky="w", pady=4)
        self.to_var = tk.StringVar()
        tk.Entry(form_card, textvariable=self.to_var, **entry_style).grid(
            row=0, column=1, sticky="ew", pady=4, ipady=3)

        tk.Label(form_card, text="Cc:", width=8, anchor="w", **label_style).grid(
            row=1, column=0, sticky="w", pady=4)
        self.cc_var = tk.StringVar()
        tk.Entry(form_card, textvariable=self.cc_var, **entry_style).grid(
            row=1, column=1, sticky="ew", pady=4, ipady=3)

        tk.Label(form_card, text="Subject:", width=8, anchor="w", **label_style).grid(
            row=2, column=0, sticky="w", pady=4)
        self.subject_var = tk.StringVar()
        tk.Entry(form_card, textvariable=self.subject_var, **entry_style).grid(
            row=2, column=1, sticky="ew", pady=4, ipady=3)

        form_card.columnconfigure(1, weight=1)

        # --- Body preview card (read-only -- no editing) ---
        body_header = tk.Frame(main, bg=COLOR_ACCENT)
        body_header.pack(fill="x")
        tk.Label(
            body_header, text="\U0001F441  Email Preview  (not editable)",
            bg=COLOR_ACCENT, fg="white", font=("Segoe UI", 10, "bold"),
            anchor="w", padx=10, pady=6,
        ).pack(fill="x")

        # --- Buttons (bottom) -- packed BEFORE body_card so this bar always
        # keeps its reserved space regardless of how tall the preview content
        # is; body_card fills what's left and scrolls internally instead of
        # pushing the buttons off-screen.
        btn_frame = tk.Frame(main, bg=COLOR_BG)
        btn_frame.pack(fill="x", side="bottom", pady=(12, 0))
        tk.Button(
            btn_frame, text="\u2713  Send", width=18, bg=COLOR_SEND_BG, fg="white",
            font=("Segoe UI", 11, "bold"), relief="flat", activebackground="#1e8449",
            command=self._on_send,
        ).pack(side="left", padx=(0, 10), ipady=6)
        tk.Button(
            btn_frame, text="\u2716  Cancel Sending", width=18, bg=COLOR_CANCEL_BG, fg="white",
            font=("Segoe UI", 11, "bold"), relief="flat", activebackground="#ba6a12",
            command=self._on_cancel,
        ).pack(side="left", ipady=6)

        body_card = tk.Frame(main, bg=COLOR_CARD_BG, highlightbackground=COLOR_ACCENT,
                              highlightthickness=1)
        body_card.pack(fill="both", expand=True)

        if self.preview_available:
            self.preview_frame = HtmlFrame(body_card, messages_enabled=False)
            self.preview_frame.pack(fill="both", expand=True, padx=2, pady=2)
        else:
            # No preview library -- show a disabled (read-only) plain-text view instead.
            hint = tk.Label(
                body_card,
                text="(Install 'tkinterweb' -- pip install tkinterweb -- for a "
                     "fully rendered preview. Showing plain text below.)",
                bg="#fff8e1", fg="#8a6100", font=("Segoe UI", 9, "italic"),
                anchor="w", padx=8, pady=4,
            )
            hint.pack(fill="x")
            text_frame = tk.Frame(body_card, bg=COLOR_CARD_BG)
            text_frame.pack(fill="both", expand=True)
            scrollbar = tk.Scrollbar(text_frame)
            scrollbar.pack(side="right", fill="y")
            self.preview_text = tk.Text(
                text_frame, wrap="word", font=("Consolas", 10),
                yscrollcommand=scrollbar.set, state="disabled",
                bg="#fbfcfd", relief="flat",
            )
            self.preview_text.pack(side="left", fill="both", expand=True, padx=6, pady=6)
            scrollbar.config(command=self.preview_text.yview)
            self.preview_frame = None

        tk.Button(
            side, text="\u23F9\nEnd", width=8, bg=COLOR_END_BG, fg="white",
            font=("Segoe UI", 10, "bold"), relief="flat", activebackground="#922b21",
            command=self._on_end,
        ).pack(pady=4, ipady=10)

    # -- flow ---------------------------------------------------------------
    def _load_current(self):
        if self.index >= len(self.branches):
            self._show_summary()
            return
        b = self.branches[self.index]
        default_to = self.email_map.get(b["branch_code"], "")
        subject = SUBJECT_TEMPLATE.format(branch_name=b["branch_name"], branch_code=b["branch_code"])
        self.current_body = build_default_body(b["rows"])

        self.to_var.set(default_to)
        self.cc_var.set(DEFAULT_CC)
        self.subject_var.set(subject)
        self._render_body_preview()

        self.progress_label.config(
            text=f"Email {self.index + 1} of {len(self.branches)}  —  "
                 f"{b['branch_code']} ({b['branch_name']})"
        )

    def _render_body_preview(self):
        if self.preview_available and self.preview_frame is not None:
            try:
                self.preview_frame.load_html(self.current_body)
            except Exception:
                pass
        elif hasattr(self, "preview_text"):
            plain = html_to_plain_fallback(self.current_body)
            self.preview_text.config(state="normal")
            self.preview_text.delete("1.0", "end")
            self.preview_text.insert("1.0", plain)
            self.preview_text.config(state="disabled")

    def _current_fields(self):
        return {
            "to": self.to_var.get().strip(),
            "cc": self.cc_var.get().strip(),
            "subject": self.subject_var.get().strip(),
            "body": self.current_body,
        }

    def _on_send(self):
        from tkinter import messagebox
        b = self.branches[self.index]
        fields = self._current_fields()
        to_list = parse_address_list(fields["to"])
        cc_list = parse_address_list(fields["cc"])

        if not to_list:
            messagebox.showwarning("Missing recipient", "Please enter at least one address in To before sending.")
            return

        try:
            if self.smtp_conn is None:
                self.smtp_conn = connect_smtp()
            msg = EmailMessage()
            msg["From"] = self.sender_email
            msg["To"] = ", ".join(to_list)
            if cc_list:
                msg["Cc"] = ", ".join(cc_list)
            msg["Subject"] = fields["subject"]
            msg["Date"] = formatdate(localtime=True)
            msg["Message-ID"] = make_msgid()
            msg.set_content(html_to_plain_fallback(fields["body"]))
            msg.add_alternative(fields["body"], subtype="html")
            self.smtp_conn.send_message(msg)
        except Exception as e:
            messagebox.showerror("Send failed", f"Could not send this email:\n{e}\n\n"
                                                 f"You can retry, edit To/Cc/Subject, or click Cancel Sending.")
            return

        self.results.append({
            "branch_code": b["branch_code"], "branch_name": b["branch_name"],
            "status": "sent",
            "to": ", ".join(to_list), "cc": ", ".join(cc_list),
            "subject": fields["subject"], "body": fields["body"], "rows": b["rows"],
        })
        self.sent_count += 1
        self.index += 1
        self._load_current()

    def _on_cancel(self):
        b = self.branches[self.index]
        fields = self._current_fields()
        self.results.append({
            "branch_code": b["branch_code"], "branch_name": b["branch_name"],
            "status": "not_sent",
            "to": fields["to"], "cc": fields["cc"],
            "subject": fields["subject"], "body": "",
        })
        self.not_sent_count += 1
        self.index += 1
        self._load_current()

    def _on_end(self):
        for b in self.branches[self.index:]:
            self.results.append({
                "branch_code": b["branch_code"], "branch_name": b["branch_name"],
                "status": "not_sent", "to": "", "cc": "", "subject": "", "body": "",
            })
        self.not_sent_count += len(self.branches) - self.index
        self.index = len(self.branches)
        self._show_summary()

    def _show_summary(self):
        import tkinter as tk
        for w in self.root.winfo_children():
            w.destroy()
        self.root.configure(bg=COLOR_HEADER_BG)
        frame = tk.Frame(self.root, bg=COLOR_HEADER_BG)
        frame.pack(expand=True)
        tk.Label(
            frame, text="\u2705 Process Completed", font=("Segoe UI", 22, "bold"),
            bg=COLOR_HEADER_BG, fg="white",
        ).pack(pady=(40, 16))
        tk.Label(
            frame, text=f"Emails sent: {self.sent_count}", font=("Segoe UI", 14, "bold"),
            bg=COLOR_HEADER_BG, fg=COLOR_SEND_BG,
        ).pack(pady=4)
        tk.Label(
            frame, text=f"Emails not sent: {self.not_sent_count}", font=("Segoe UI", 14, "bold"),
            bg=COLOR_HEADER_BG, fg=COLOR_CANCEL_BG,
        ).pack(pady=4)
        tk.Button(
            frame, text="Close", width=16, bg=COLOR_ACCENT, fg="white",
            font=("Segoe UI", 11, "bold"), relief="flat", command=self.root.destroy,
        ).pack(pady=28, ipady=6)

# ==========================================================================
# Word summary document
# ==========================================================================

def add_word_table(document, rows):
    """Add a native Word table for one branch's account rows (mirrors what
    was actually shown/sent in the HTML email)."""
    from docx.shared import RGBColor
    headings = [h for _, h in COLUMN_MAP]
    table = document.add_table(rows=1, cols=len(headings))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(headings):
        hdr_cells[i].text = heading
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for record in rows:
        cells = table.add_row().cells
        for i, heading in enumerate(headings):
            cells[i].text = record.get(heading, "")
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    return table


def build_summary_document(results, output_path):
    document = Document()
    document.add_heading("Branch Activation Emails — SMIFS Limited", level=1)

    sent = [r for r in results if r["status"] == "sent"]
    not_sent = [r for r in results if r["status"] == "not_sent"]

    p = document.add_paragraph()
    p.add_run(f"Run date: {datetime.now().strftime('%d-%m-%Y %H:%M')}").italic = True

    document.add_heading(f"Emails Sent ({len(sent)})", level=2)
    if not sent:
        document.add_paragraph("None.")
    for r in sent:
        document.add_heading(f"{r['branch_code']} — {r['branch_name']}", level=3)
        document.add_paragraph(f"To: {r['to']}")
        if r["cc"]:
            document.add_paragraph(f"Cc: {r['cc']}")
        document.add_paragraph(f"Subject: {r['subject']}")

        document.add_paragraph(EMAIL_GREETING)
        document.add_paragraph(EMAIL_INTRO_1)
        document.add_paragraph(EMAIL_INTRO_2)

        if r.get("rows"):
            add_word_table(document, r["rows"])
        else:
            # Fallback for any older/edited body without row data attached.
            document.add_paragraph(html_to_plain_fallback(r["body"]))

        document.add_paragraph()
        nb_para = document.add_paragraph()
        nb_para.add_run(EMAIL_NB).bold = True
        document.add_paragraph(EMAIL_HELPDESK)
        for line in EMAIL_SIGNATURE_LINES:
            document.add_paragraph(line).runs[0].italic = True

        document.add_paragraph()  # spacer between emails

    document.add_heading(f"Emails Not Sent ({len(not_sent)})", level=2)
    if not not_sent:
        document.add_paragraph("None.")
    else:
        document.add_paragraph("In the sequence they were reached:")
        for r in not_sent:
            document.add_paragraph(f"{r['branch_code']} — {r['branch_name']}", style="List Bullet")

    document.save(output_path)


def with_date_suffix(path):
    """Insert the current system date before the file extension."""
    base, ext = os.path.splitext(path)
    date_str = datetime.now().strftime("%d-%m-%Y")
    return f"{base}_{date_str}{ext}"


# ==========================================================================
# File pickers (input files + output location)
# ==========================================================================

def _gui_available():
    try:
        import tkinter as tk
        root = tk.Tk()
        root.withdraw()
        root.destroy()
        return True
    except Exception:
        return False


def ask_open_file(label, default_filename):
    if _gui_available():
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        path = filedialog.askopenfilename(
            title=f"Select {label}",
            filetypes=[("Excel files", "*.xlsx *.xlsm"), ("All files", "*.*")],
        )
        root.destroy()
        if not path:
            print(f"No file selected for {label}. Exiting.")
            sys.exit(1)
        return path

    default_path = os.path.join(os.getcwd(), default_filename)
    entered = input(f"Path to {label} [{default_path}]: ").strip()
    return entered if entered else default_path


def main():
    print("=== Branch Activation Email Reviewer ===")
    ao_path = ask_open_file("AOsheet.xlsx", "AOsheet.xlsx")
    email_path = ask_open_file("EmailIDs.xlsx", "EmailIDs.xlsx")

    for p in (ao_path, email_path):
        if not os.path.isfile(p):
            print(f"ERROR: file not found: {p}")
            sys.exit(1)

    print("Reading AOsheet.xlsx and grouping by Branch Code (column H)...")
    groups = read_ao_sheet(ao_path)
    print(f"  Found {len(groups)} distinct branch codes.")

    print("Reading EmailIDs.xlsx (Column A -> Column C)...")
    email_map = read_email_map(email_path)
    print(f"  Loaded {len(email_map)} branch->email mappings.")

    branches = [
        {"branch_code": code, "branch_name": data["branch_name"], "rows": data["rows"]}
        for code, data in groups.items()
    ]

    import tkinter as tk
    root = tk.Tk()
    app = EmailReviewApp(root, branches, email_map, SENDER_EMAIL)
    root.mainloop()

    # Output saved automatically next to AOsheet.xlsx -- no save dialog.
    output_dir = os.path.dirname(os.path.abspath(ao_path))
    output_path = with_date_suffix(os.path.join(output_dir, "Branch_Activation_Emails.docx"))
    build_summary_document(app.results, output_path)

    print()
    print(f"Emails sent: {app.sent_count}, not sent: {app.not_sent_count}.")
    print(f"Summary Word document saved to: {output_path}")
    print("Process completed")


if __name__ == "__main__":
    main()